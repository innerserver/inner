from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "Inner_App_User_Manual.docx"


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "Arial"
    run.font.size = Pt(10)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP


def add_heading(doc, text, level=1):
    paragraph = doc.add_heading(text, level=level)
    for run in paragraph.runs:
      run.font.name = "Arial"
      run.font.color.rgb = RGBColor(21, 21, 21)
    return paragraph


def add_body(doc, text):
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(10.5)
    return paragraph


def add_bullets(doc, items):
    for item in items:
        paragraph = doc.add_paragraph(style="List Bullet")
        run = paragraph.add_run(item)
        run.font.name = "Arial"
        run.font.size = Pt(10.5)


def add_numbers(doc, items):
    for item in items:
        paragraph = doc.add_paragraph(style="List Number")
        run = paragraph.add_run(item)
        run.font.name = "Arial"
        run.font.size = Pt(10.5)


def build():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(10.5)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("Inner App User Manual")
    title_run.bold = True
    title_run.font.name = "Arial"
    title_run.font.size = Pt(24)
    title_run.font.color.rgb = RGBColor(15, 118, 110)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.add_run("Tutorials for members, moderators, admins, and HMD/dev users")
    subtitle_run.font.name = "Arial"
    subtitle_run.font.size = Pt(11)
    subtitle_run.font.color.rgb = RGBColor(98, 100, 95)

    add_heading(doc, "Start Here", 1)
    add_body(doc, "Inner is a private workspace app with accounts, rooms, DMs, files, browser tools, Google Docs, Google Slides, Google Sheets, Admin controls, and HMD/dev tools.")
    add_bullets(doc, [
        "The Dashboard has a role-based guide after login.",
        "The HTML manual is available in the app at /tutorials.html.",
        "Normal users should not see Admin, HMD, or Domain tabs.",
        "Inner Docs stays experimental while Google Docs, Slides, and Sheets are the main work tools.",
    ])

    add_heading(doc, "Creating or Requesting an Account", 1)
    add_numbers(doc, [
        "Open Inner and press Create or request account.",
        "When the browser asks for location, choose Always allow.",
        "Fill in real display name, email, phone, grade, username, and password.",
        "Open signup creates a normal member account.",
        "Request mode sends the account request to admins and shows a review screen until a decision is made.",
    ])

    add_heading(doc, "Tutorials by Account Type", 1)
    role_table = doc.add_table(rows=1, cols=3)
    role_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    role_table.style = "Table Grid"
    headers = role_table.rows[0].cells
    for cell, text in zip(headers, ["Account type", "Main actions", "Notes"]):
        shade_cell(cell, "EDF7F4")
        set_cell_text(cell, text, True)
    rows = [
        ("Member / Student", "Set profile, find friends, use messages/DMs, upload files, use Google Docs/Slides/Sheets, share links.", "Can friend same-grade users normally or search exact username/email/phone."),
        ("Moderator / Teacher", "Use Messages, DMs, docs, and reports/mod tools when granted.", "Moderator accounts should be assigned manually by admins."),
        ("Admin", "Manage signup, accounts, announcements, hidden tabs, locks, paywalls, browser rules, logs, reports, and backups.", "Search accounts first; use Show all only when needed."),
        ("HMD / Dev", "Review health, storage, database, logs, localhost tools, bots, plugins, AI, and emergency controls.", "Use shutdown/recovery tools carefully."),
    ]
    for row in rows:
        cells = role_table.add_row().cells
        for cell, text in zip(cells, row):
            set_cell_text(cell, text)

    add_heading(doc, "Docs, Slides, Sheets, and Browser", 1)
    docs_table = doc.add_table(rows=1, cols=3)
    docs_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    docs_table.style = "Table Grid"
    headers = docs_table.rows[0].cells
    for cell, text in zip(headers, ["Area", "What it does", "Important note"]):
        shade_cell(cell, "EDF7F4")
        set_cell_text(cell, text, True)
    rows = [
        ("Google Docs", "Create and edit documents inside Inner.", "If Google blocks sign-in in the frame, open full tab once, then return."),
        ("Google Slides", "Create and edit presentations inside Inner.", "Files save in the user's Google account."),
        ("Google Sheets", "Create and edit spreadsheets inside Inner.", "Works like the Chess embedded app tab."),
        ("Inner Docs", "Experimental built-in docs, slides, and notes saved inside Inner.", "Use for testing until fully upgraded."),
        ("Browser", "Open allowed sites inside Inner and send links to friends/groups.", "Admin block/allow rules apply through Inner's browser route."),
    ]
    for row in rows:
        cells = docs_table.add_row().cells
        for cell, text in zip(cells, row):
            set_cell_text(cell, text)

    add_heading(doc, "Admin Checklist", 1)
    add_bullets(doc, [
        "Set Admin > Server > Signup mode.",
        "Add report email recipients.",
        "Hide Admin, HMD, Domain, or any feature from normal users as needed.",
        "Use Account search by name, username, email, phone, grade, IP, device, or role.",
        "Use Browser rules to block or allow websites for normal users.",
        "Export logs before wiping them.",
        "Create backups before major changes and restore from backups if needed.",
    ])

    add_heading(doc, "Phone Setup", 1)
    add_heading(doc, "iPhone", 2)
    add_numbers(doc, [
        "Open Inner in Safari.",
        "Tap Share.",
        "Tap Add to Home Screen.",
        "Open Inner from the home-screen icon.",
        "Use the three-line menu to open the sidebar.",
    ])
    add_heading(doc, "Android", 2)
    add_numbers(doc, [
        "Open Inner in Chrome.",
        "Tap Install app or Add to Home screen.",
        "Open Inner from the installed icon.",
        "Use the three-line menu to open the sidebar.",
    ])

    footer = section.footer.paragraphs[0]
    footer.text = "Inner Manual - keep with the deployed app and /tutorials.html"
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
